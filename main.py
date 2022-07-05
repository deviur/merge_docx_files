import docx
import os

"""
TODO: write notes here ⤵
Файл doc покетом docx открыть нельзя.
"""
# doc = docx.Document()
# doc.add_paragraph('Это первый абзац')
# doc.save('example.docx')
# doc = docx.Document('example.docx')
# doc.add_paragraph('Это второй абзац')
# doc.save('example2.docx')


def get_docx_files():
    return [file for file in os.listdir() if file.endswith('.docx')]


def load_docx_files():
    with open('context.txt', 'r') as fp:
        return [line.rstrip() for line in fp.readlines()]


def append(from_where: docx.Document, to_where: docx.Document()):
    for el in from_where.element.body:
        to_where.element.body.append(el)
    return to_where


def main():

    if os.path.isfile('merged.docx'):
        os.remove('merged.docx')

    try:
        open('context.txt', 'r')
    except FileNotFoundError:
        with open('context.txt', 'w') as fp:
            files = get_docx_files()
            for f in get_docx_files():
                fp.write(f + '\n')
        return

    docx_files = load_docx_files()
    merged_document = docx.Document()
    for file in docx_files:
        sub_document = docx.Document(file)
        append(sub_document, merged_document)
    merged_document.save('merged.docx')


if __name__ == '__main__':
    main()
    # to_doc = docx.Document()
    #
    # from_doc = docx.Document('example.docx')
    # append(from_doc, to_doc)
    #
    # from_doc = docx.Document('example2.docx')
    # append(from_doc, to_doc)
    #
    # to_doc.save('example3.docx')
    # print(get_docx_files())
