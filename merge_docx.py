
import docx
import os


def get_docx_files():
    return [file for file in os.listdir() if file.endswith('.docx')]


def load_docx_files():
    with open('content_docx.txt', 'r') as fp:
        return [line.rstrip() for line in fp.readlines()]


def copy_paragraph(from_par, to_par):
    for r in from_par.runs:
        if r.style:
            to_r = to_par.add_run(r.text, r.style.name)
        else:
            to_r = to_par.add_run(r.text)

        to_r.bold = r.bold
        to_r.underline = r.underline
        to_r.italic = r.italic


def append(from_where: docx.Document, to_where: docx.Document()):
    for p in from_where.paragraphs:
        try:
            to_p = to_where.add_paragraph('', p.style.name)
        except KeyError:  # Document 'to_where' has not the style 'p.style.name'
            # TODO: Add 'p.style.name' style to the document 'to_where'
            to_p = to_where.add_paragraph('')

        to_p.alignment = p.alignment
        copy_paragraph(p, to_p)


def main():

    if os.path.isfile('merged.docx'):
        os.remove('merged.docx')

    try:
        open('content_docx.txt', 'r')
    except FileNotFoundError:
        with open('content_docx.txt', 'w') as fp:
            for f in get_docx_files():
                fp.write(f + '\n')
        return

    docx_files = load_docx_files()
    merged_document = docx.Document()
    for file in docx_files:
        sub_document = docx.Document(file)
        append(sub_document, merged_document)
    merged_document.save('merged.docx')


if __name__ == '__main__':
    main()
